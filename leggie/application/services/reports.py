"""Report Renderers — generate analysis reports.

Template Method pattern: fixed skeleton with varying steps per report type.
Builder pattern: assemble complex reports step by step.

Phase 4: 2 report types — Executive Summary + Article-by-Article.
"""

from __future__ import annotations

import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from leggie.application.agents.improver import Suggestion
from leggie.domain.models import Document, Finding, Severity


def _add_formatted_paragraph(doc: Any, text: str, style: str | None = None) -> None:
    """Add a paragraph to *doc* with basic Markdown bold/italic runs."""
    paragraph = doc.add_paragraph(style=style)
    pattern = re.compile(r"(\*\*[^*]+?\*\*|(?<!\w)_(\w[\w ]*?)_(?!\w)|\*[^*]+?\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        marker = match.group()
        run = paragraph.add_run(marker.strip("*_"))
        if marker.startswith("**"):
            run.bold = True
        else:
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


@dataclass
class Report:
    """A complete analysis report."""

    title: str
    report_type: str
    sections: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_markdown(self) -> str:
        """Render the report as Markdown."""
        lines = [f"# {self.title}", ""]
        if self.metadata:
            for k, v in self.metadata.items():
                lines.append(f"- **{k}:** {v}")
            lines.append("")

        for section in self.sections:
            level = section.get("level", 2)
            prefix = "#" * level
            lines.append(f"{prefix} {section['title']}")
            lines.append("")
            content = section.get("content", "")
            if isinstance(content, str):
                lines.append(content)
            elif isinstance(content, list):
                lines.extend(content)
            lines.append("")

        return "\n".join(lines)

    def to_docx(self, path: str | Path) -> Path:
        """Render the report as a Word document and save it to *path*."""
        from docx import Document

        doc = Document()
        doc.add_heading(self.title, level=1)

        if self.metadata:
            for key, value in self.metadata.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph()

        for section in self.sections:
            level = section.get("level", 2)
            doc.add_heading(section["title"], level=level)
            content = section.get("content", "")
            if isinstance(content, str):
                _add_formatted_paragraph(doc, content)
            elif isinstance(content, list):
                for item in content:
                    text = item.strip()
                    style = None
                    if text.startswith(("- ", "* ")):
                        text = text[2:]
                        style = "List Bullet"
                    _add_formatted_paragraph(doc, text, style=style)

        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path


class ReportRenderer(ABC):
    """Base report renderer — Template Method pattern.

    Fixed skeleton: build_header → build_body → build_footer.
    """

    @abstractmethod
    def report_type(self) -> str: ...

    async def render(
        self,
        document: Document,
        findings: list[Finding],
        suggestions: list[Suggestion] | None = None,
    ) -> Report:
        """Template Method: render the full report."""
        title = self._build_title(document)
        metadata = self._build_metadata(document, findings)
        sections = await self._build_body(document, findings, suggestions or [])
        return Report(
            title=title,
            report_type=self.report_type(),
            sections=sections,
            metadata=metadata,
        )

    def _build_title(self, document: Document) -> str:
        return f"Legal Analysis: {document.title}"

    def _build_metadata(self, document: Document, findings: list[Finding]) -> dict[str, Any]:
        by_severity = self._count_by_severity(findings)
        by_type = self._count_by_type(findings)
        return {
            "Document": document.title or "Untitled",
            "Articles analyzed": str(len(document.articles)),
            "Total findings": str(len(findings)),
            "Severity breakdown": ", ".join(f"{k}: {v}" for k, v in sorted(by_severity.items())),
            "Type breakdown": ", ".join(f"{k}: {v}" for k, v in sorted(by_type.items())),
        }

    @abstractmethod
    async def _build_body(
        self,
        document: Document,
        findings: list[Finding],
        suggestions: list[Suggestion],
    ) -> list[dict[str, Any]]: ...

    def _count_by_severity(self, findings: list[Finding]) -> dict[str, int]:
        counts: dict[str, int] = {}
        for f in findings:
            k = f.severity.value
            counts[k] = counts.get(k, 0) + 1
        return counts

    def _count_by_type(self, findings: list[Finding]) -> dict[str, int]:
        counts: dict[str, int] = {}
        for f in findings:
            k = f.finding_type.value
            counts[k] = counts.get(k, 0) + 1
        return counts

    def _findings_by_article(self, findings: list[Finding]) -> dict[str, list[Finding]]:
        by_article: dict[str, list[Finding]] = {}
        for f in findings:
            # Extract article ID from IRAC issue text
            parts = f.irac.issue.split(" ")
            article_id = ""
            for p in parts:
                if p.isdigit():
                    article_id = p
                    break
            if not article_id:
                article_id = "general"
            if article_id not in by_article:
                by_article[article_id] = []
            by_article[article_id].append(f)
        return by_article


class ExecutiveSummaryRenderer(ReportRenderer):
    """Executive Summary — high-level overview of findings."""

    def report_type(self) -> str:
        return "executive_summary"

    async def _build_body(
        self,
        document: Document,
        findings: list[Finding],
        suggestions: list[Suggestion],
    ) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []

        # Overview
        sections.append(
            {
                "level": 2,
                "title": "Overview",
                "content": (
                    f"This report analyzes **{document.title}** ({len(document.articles)} articles, "
                    f"{len(findings)} findings). "
                    f"The analysis identified {sum(1 for f in findings if f.severity in (Severity.CRITICAL, Severity.HIGH))} "
                    f"high-severity issues."
                ),
            }
        )

        # Findings summary by severity
        critical = [f for f in findings if f.severity.value == "critical"]
        high = [f for f in findings if f.severity.value == "high"]
        medium = [f for f in findings if f.severity.value == "medium"]

        if critical:
            sections.append(
                {
                    "level": 2,
                    "title": "Critical Issues",
                    "content": [
                        f"- **{f.irac.issue[:100]}** — {f.finding_type.value} ({f.lens})"
                        for f in critical
                    ],
                }
            )
        if high:
            sections.append(
                {
                    "level": 2,
                    "title": "High-Severity Issues",
                    "content": [
                        f"- **{f.irac.issue[:100]}** — {f.finding_type.value} ({f.lens})"
                        for f in high
                    ],
                }
            )
        if medium:
            sections.append(
                {
                    "level": 2,
                    "title": "Medium-Severity Issues",
                    "content": [f"- {f.irac.issue[:100]} ({f.lens})" for f in medium],
                }
            )

        # Suggestions
        if suggestions:
            sections.append(
                {
                    "level": 2,
                    "title": "Recommendations",
                    "content": [
                        f"- {s.description} (priority: {s.priority})" for s in suggestions[:10]
                    ],
                }
            )

        return sections


class ArticleByArticleRenderer(ReportRenderer):
    """Article-by-Article — detailed findings per article."""

    def report_type(self) -> str:
        return "article_by_article"

    async def _build_body(
        self,
        document: Document,
        findings: list[Finding],
        suggestions: list[Suggestion],
    ) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        by_article = self._findings_by_article(findings)

        for article_id in sorted(
            by_article.keys(), key=lambda x: (not x.isdigit(), int(x) if x.isdigit() else 0)
        ):
            article_findings = by_article[article_id]
            article = None
            for a in document.articles:
                if a.id == article_id:
                    article = a
                    break

            header_title = f"Άρθρο {article_id}"
            if article and article.title:
                header_title += f": {article.title}"

            content_lines: list[str] = []
            for f in article_findings:
                content_lines.append(
                    f"**{f.finding_type.value.upper()}** — Severity: {f.severity.value}, Confidence: {f.confidence.score}"
                )
                content_lines.append(f"  - Issue: {f.irac.issue}")
                content_lines.append(f"  - Conclusion: {f.irac.conclusion}")
                if f.evidence:
                    for e in f.evidence:
                        if e.text_excerpt:
                            content_lines.append(f"  - Evidence: _{e.text_excerpt}_")
                content_lines.append("")

            # Article-specific suggestions
            article_suggestions = [s for s in suggestions if s.article_id == article_id]
            if article_suggestions:
                content_lines.append("**Suggestions:**")
                for s in article_suggestions:
                    content_lines.append(f"- [{s.suggestion_type}] {s.description}")
                content_lines.append("")

            sections.append(
                {
                    "level": 2,
                    "title": header_title,
                    "content": content_lines,
                }
            )

        # General suggestions (no specific article)
        general = [s for s in suggestions if not s.article_id]
        if general:
            sections.append(
                {
                    "level": 2,
                    "title": "General Recommendations",
                    "content": [f"- [{s.suggestion_type}] {s.description}" for s in general],
                }
            )

        return sections
