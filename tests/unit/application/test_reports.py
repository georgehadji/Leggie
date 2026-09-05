"""Tests for report renderers."""

import pytest
from docx import Document as DocxDocument

from leggie.application.services.reports import ArticleByArticleRenderer, Report
from leggie.domain.models import (
    IRAC,
    Article,
    Confidence,
    Document,
    Finding,
    FindingType,
    Severity,
)


def _finding(article_id: str, issue: str) -> Finding:
    return Finding(
        finding_type=FindingType.CONSTITUTIONAL,
        article_id=article_id,
        irac=IRAC(issue=issue, rule="r", application="a", conclusion="c"),
        severity=Severity.HIGH,
        confidence=Confidence.from_score(0.7),
        lens="test",
        model="test",
    )


def _document(article_ids: list[str]) -> Document:
    return Document(
        title="Test Bill",
        source_format="txt",
        articles=[Article(id=aid, title=f"Title {aid}", raw_text="text") for aid in article_ids],
    )


class TestFindingsByArticleGrouping:
    """DH-14: _findings_by_article used to hand-parse irac.issue text
    (split on space, look for a bare-digit token) instead of using the
    reliable, always-populated finding.article_id field via
    article_number_of() — the exact helper D1/DH-13 already established
    for this same failure mode elsewhere (cove_verifier.py,
    blackboard_aggregator.py, bill_analysis_flow.py, agents/improver.py).

    Every lens's own regex-fallback issue text has the shape
    "Άρθρο {id}: ...", colon immediately after the digit, so `"5:".isdigit()`
    is False — the naive parse found ZERO article id for every such finding
    and silently dumped the entire Article-by-Article report (one of only
    two report types this system produces) into a single "general" section.
    LLM-authored issue text has no guaranteed digit token at all, so the old
    code was broken for both the regex-fallback path and the common
    LLM-generated path — this had zero pre-existing test coverage.
    """

    @pytest.mark.asyncio
    async def test_regex_fallback_issue_shape_groups_by_real_article(self):
        # Exactly the shape every lens's own regex-fallback path produces.
        f = _finding("5", "Άρθρο 5: Πιθανή υπέρβαση ορίων νομοθετικής εξουσιοδότησης")
        renderer = ArticleByArticleRenderer()

        report = await renderer.render(_document(["5"]), [f])

        titles = [s["title"] for s in report.sections]
        assert titles == ["Άρθρο 5: Title 5"]

    @pytest.mark.asyncio
    async def test_llm_freeform_issue_text_still_groups_by_structured_field(self):
        # LLM-authored issue text with no digit token at all — the naive
        # split-and-isdigit parse would always miss this too.
        f = _finding("12", "Δεν αναφέρεται ρητά ο τρόπος υπολογισμού των τελών")
        renderer = ArticleByArticleRenderer()

        report = await renderer.render(_document(["12"]), [f])

        titles = [s["title"] for s in report.sections]
        assert titles == ["Άρθρο 12: Title 12"]

    @pytest.mark.asyncio
    async def test_two_articles_are_not_collapsed_into_one_section(self):
        findings = [
            _finding("1", "Άρθρο 1: πρώτο ζήτημα"),
            _finding("2", "Άρθρο 2: δεύτερο ζήτημα"),
        ]
        renderer = ArticleByArticleRenderer()

        report = await renderer.render(_document(["1", "2"]), findings)

        titles = {s["title"] for s in report.sections}
        assert titles == {"Άρθρο 1: Title 1", "Άρθρο 2: Title 2"}

    @pytest.mark.asyncio
    async def test_legacy_finding_with_no_article_id_and_no_number_falls_to_general(self):
        """No-regression: a genuinely legacy/unattributable finding still
        degrades to the "general" bucket, not a crash or a fabricated id."""
        f = _finding("", "Γενικό ζήτημα χωρίς αναφορά άρθρου")
        renderer = ArticleByArticleRenderer()

        report = await renderer.render(_document([]), [f])

        titles = [s["title"] for s in report.sections]
        assert titles == ["Άρθρο general"]


def test_report_to_docx_renders_bold_and_italic(tmp_path):
    """DOCX output preserves basic Markdown bold/italic formatting."""
    report = Report(
        title="Test Report",
        report_type="test",
        metadata={},
        sections=[
            {
                "level": 2,
                "title": "Formatted section",
                "content": "This is **bold** and this is _italic_ text.",
            },
            {
                "level": 2,
                "title": "Bullets",
                "content": [
                    "- **bold** bullet",
                    "- _italic_ bullet",
                    "- plain bullet",
                ],
            },
        ],
    )

    path = tmp_path / "report.docx"
    result = report.to_docx(path)

    assert result == path
    assert path.exists()

    doc = DocxDocument(str(path))
    bold_runs = [run.text for p in doc.paragraphs for run in p.runs if run.bold]
    italic_runs = [run.text for p in doc.paragraphs for run in p.runs if run.italic]

    assert "bold" in bold_runs
    assert "italic" in italic_runs
